#! /usr/bin/env python
# -*- coding: utf-8 -*-

import xml.etree.ElementTree as ET
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

tree = ET.parse('12.04.02-01-00-ИБ-ver7.plm.xml')
root = tree.getroot()
competences=root.find('План').find('Компетенции')
comp=[]
for child in competences:
    comp.append([child.get('Индекс'),child.get('Содержание')])
    #print(child.get('Индекс'))
    #print(child.get('Содержание'))

subjects=root.find('План').find('СтрокиПлана')

theDirection='Оптотехника'
theCodeOfDirection='12.04.02'
who='магистр'
whom='магистров'
author='Костюченко Владимир Яковлевич, д-р физ.-мат. наук доцент, профессор каф. Физики'
authorComments='Ф.И.О., степень, звание'
theDate='28.11.2014 г.'


def getCContent(competences,c):
    for x in competences:
        if x[0]==c:
            return x[1]

subj=[]
for child in subjects:

    subj.append([child.get('Дис'), # 0
                 child.get('ИдетификаторДисциплины'), # 1
                child.get('СемЗач'), # 2
                 child.get('КредитовНаДисциплину'), # 3
                child.find('Сем').get('Зач'), # 4
                child.get('ПодлежитИзучению'), # 5
                child.find('Сем').get('Лек'), # 6
                child.find('Сем').get('Пр'), # 7
                child.find('Сем').get('СРС'), # 8
                child.find('Сем').get('ЧасЭкз'), #9
                child.find('Сем').get('Экз'), # 10
                child.get('Компетенции'), # 11
                child.get('СемЭкз') # 12
                ]) 

for s in range(0,len(subj)):
    
    if(subj[s][11]!=None):
        subj[s][11]=subj[s][11].split(', ')

        
    #print(child.get('Дис')) # Имя дисциплины
    #print(child.get('ИдетификаторДисциплины')) # Цикл, раздел учебного плана
    #print(child.get('СемЗач')) # Курс изучения
    #print(child.get('КредитовНаДисциплину')) # Количество зачетных единиц
    #print(child.find('Сем').get('Зач')) # Форма промежуточной аттестации
    #print(child.get('ПодлежитИзучению')) # Количество часов всего
    #print(child.find('Сем').get('Лек')) # Лекционные часы
    #print(child.find('Сем').get('Пр')) # Практические
    # print(child.get('Дис')) # Лабораторные из везде 0
    #print(child.find('Сем').get('СРС')) # СРС
    #print(child.get('Дис')) # подготовка к экзамену
     # Цели
    #print(child.get('Компетенции')) # Компетенции
     # Краткое содержание дисциплины


for s in subj:

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    paragraph = document.add_paragraph('АННОТАЦИЯ',style='Normal')

    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph('к рабочей программе дисциплины')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph('«'+s[0] +'»')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph('                       Составитель:')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph('«'+author +'»')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph = document.add_paragraph('«'+authorComments +'»')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = document.add_table(rows=14, cols=2)
    row = table.columns[0]
    row.cells[0].text = 'Направление подготовки'
    row.cells[1].text = 'Профиль подготовки'
    row.cells[2].text = 'Квалификация (степень) выпускника'
    row.cells[3].text = 'Форма обучения'
    row.cells[4].text = 'Цикл, раздел учебного плана'
    row.cells[5].text = 'Курс изучения'
    row.cells[6].text = 'Количество зачетных единиц'
    row.cells[7].text = 'Форма промежуточной аттестации'
    row.cells[8].text = 'Количество часов всего, из них'
    row.cells[9].text = '- лекционные'
    row.cells[10].text = '- практические'
    row.cells[11].text = '- лабораторные'
    row.cells[12].text = '- СРС'
    row.cells[13].text = '- подготовка к экзамену'

    row = table.columns[1]
    row.cells[0].text = theCodeOfDirection+' '+theDirection
    row.cells[1].text = 'Компьютерная безопасность'
    row.cells[2].text = who
    row.cells[3].text = 'очная'
    row.cells[4].text = 'Вариативная часть ' + s[1]
    if s[2]!=None:
        row.cells[5].text = str(int(s[2]) / 2)
    else if s[12]!=None:
        row.cells[5].text = str(int(s[12]) / 2)
    else:
        row.cells[5].text = 'Поле не заполнено!!! ' #str(int(s[2]) / 2)

    row.cells[6].text = s[3]

    row.cells[7].text = 'Произошла ошибка обрабоки!!!'
    if s[4]!=None and s[10]==None:
        row.cells[7].text = 'Зачет.'
    else:
        row.cells[7].text = 'Экзамен.'

    row.cells[8].text = s[5]
    if s[6]!=None:
         row.cells[9].text = s[6]
    else:
         row.cells[9].text = '0'
    if s[7]!=None:
         row.cells[10].text = s[7]
    else:
         row.cells[10].text = '0'

    row.cells[11].text = '0'

    if s[8]!=None:
         row.cells[12].text = s[8]
    else:
         row.cells[12].text = '0'

    if s[9]!=None:
        row.cells[13].text = str(s[9])
    else:
        row.cells[13].text = '0'

    paragraph = document.add_paragraph('1.	Целями освоения дисциплины ' +
    	s[0]+'являются формирование у студентов общекультурных, общепрофессиональных и профессиональных компетенций, определяющих их готовность и способность, как будущих специалистов по направлению подготовки «'+
    	theDirection+'», к эффективному применению усвоенных знаний для'+
    	'РЕДАКТИРУЙ ЗДЕСЬ!')
    paragraph = document.add_paragraph('2.	Компетенции обучающегося, формируемые в результате освоения дисциплины:')

    okadded=False
    opkadded=False
    pkadded=False
    if s[11]!=None:
        print(s[11])
        for c in s[11]:
            print(c)
            if c[0]=='О' and c[1]!='П' and not okadded:
                paragraph = document.add_paragraph('общекультурные компетенции:')
                okadded=True
            if c[0]=='О' and c[1]!='П':
                paragraph = document.add_paragraph(c+' '+getCContent(comp,c))
            if c[0]=='О' and c[1]=='П' and not opkadded:
                paragraph = document.add_paragraph('общепрофессиональные компетенции:')
                opkadded=True
            if c[0]=='О' and c[1]=='П':
                paragraph = document.add_paragraph(c+' '+getCContent(comp,c))
            if c[0]=='П' and not pkadded:
                paragraph = document.add_paragraph('профессиональные компетенции:')
                pkadded=True
            if c[0]=='П':
                paragraph = document.add_paragraph(c+' '+getCContent(comp,c))
    else:
        paragraph = document.add_paragraph('Компетенции не заполнены!!!!')
    paragraph = document.add_paragraph('3.	Краткое содержание дисциплины ')
    paragraph = document.add_paragraph('Раздел I. Введение.')


    paragraph = document.add_paragraph('Аннотация разработана на основании ФГОС ВО по направлению подготовки '
    	+whom+' '+theCodeOfDirection+' '+theDirection + ' от '+theDate)
    

    # docx.section.Sections тут надо размер страницы нормальный сделать, A4
    # и поля.

    document.save(s[0]+'.docx')

    # os.system('touch \"//root//AnnotationGenerator//shortContent//'+s[0]+'.txt\"')
